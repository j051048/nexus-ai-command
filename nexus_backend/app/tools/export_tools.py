"""
Excel/PDF 报表导出工具
支持销售数据、分析报告的导出
"""

import io
import logging
import re
from datetime import datetime
from typing import Any, Literal

from langchain_core.tools import tool

logger = logging.getLogger(__name__)


def _clean_markdown_inline(value: str) -> str:
    value = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", value)
    value = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", value)
    return re.sub(r"[*_`~]", "", value).strip()


@tool
async def export_to_excel(
    data: list[dict[str, Any]],
    filename: str = "",
    sheet_name: str = "Sheet1",
    include_header: bool = True,
) -> dict[str, Any]:
    """导出数据到 Excel 文件

    Args:
        data: 要导出的数据列表，每个元素是一行数据的字典
        filename: 文件名（不含扩展名），为空则自动生成
        sheet_name: 工作表名称
        include_header: 是否包含表头

    Returns:
        包含文件路径或base64编码的字典

    Example:
        export_to_excel(
            data=[
                {"客户名称": "A公司", "销售额": 100000, "状态": "成交"},
                {"客户名称": "B公司", "销售额": 50000, "状态": "跟进中"}
            ],
            filename="销售数据_2026Q1"
        )
    """
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Alignment, Font, PatternFill

        if not data:
            return {"success": False, "error": "数据为空"}

        # 创建工作簿
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        # 写入表头
        headers = list(data[0].keys())
        if include_header:
            for col_idx, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col_idx, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(
                    start_color="CCCCCC", end_color="CCCCCC", fill_type="solid"
                )
                cell.alignment = Alignment(horizontal="center")

        # 写入数据
        start_row = 2 if include_header else 1
        for row_idx, row_data in enumerate(data, start_row):
            for col_idx, header in enumerate(headers, 1):
                ws.cell(row=row_idx, column=col_idx, value=row_data.get(header))

        # 自动调整列宽
        for col in ws.columns:
            max_length = 0
            column = col[0].column_letter
            for cell in col:
                if cell.value:
                    max_length = max(max_length, len(str(cell.value)))
            ws.column_dimensions[column].width = min(max_length + 2, 50)

        # 生成文件名
        if not filename:
            filename = f"export_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # 保存到内存
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        import base64

        content_base64 = base64.b64encode(output.read()).decode()

        return {
            "success": True,
            "filename": f"{filename}.xlsx",
            "content_base64": content_base64,
            "rows": len(data),
        }

    except Exception as e:
        logger.error(f"导出Excel失败: {e}")
        return {"success": False, "error": str(e)}


@tool
async def export_to_pdf(
    content: str,
    filename: str = "",
    title: str = "",
    format_type: Literal["markdown", "html"] = "markdown",
) -> dict[str, Any]:
    """导出内容到 PDF 文件

    Args:
        content: 要导出的内容（Markdown 或 HTML）
        filename: 文件名（不含扩展名）
        title: 文档标题
        format_type: 内容格式类型

    Returns:
        包含文件路径或base64编码的字典

    Example:
        export_to_pdf(
            content="# 销售分析报告\\n\\n本月销售额: 100万",
            title="2026年Q1销售分析",
            format_type="markdown"
        )
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        # 注册中文字体（使用系统字体）
        try:
            pdfmetrics.registerFont(TTFont("SimSun", "SimSun.ttf"))
            font_name = "SimSun"
        except Exception:
            font_name = "Helvetica"  # 回退到默认字体

        # 转换 Markdown 到 HTML
        if format_type == "markdown":
            import markdown

            content = markdown.markdown(content)

        # 创建 PDF
        output = io.BytesIO()
        doc = SimpleDocTemplate(
            output, pagesize=A4, topMargin=2 * cm, bottomMargin=2 * cm
        )

        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(name="Chinese", fontName=font_name, fontSize=12, leading=18)
        )

        story = []

        # 添加标题
        if title:
            title_style = ParagraphStyle(
                "Title", parent=styles["Heading1"], fontName=font_name, fontSize=18
            )
            story.append(Paragraph(title, title_style))
            story.append(Spacer(1, 0.5 * cm))

        # 添加内容（简化处理）
        for line in content.split("\n"):
            if line.strip():
                story.append(Paragraph(line, styles["Chinese"]))
                story.append(Spacer(1, 0.3 * cm))

        doc.build(story)
        output.seek(0)

        import base64

        content_base64 = base64.b64encode(output.read()).decode()

        if not filename:
            filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        return {
            "success": True,
            "filename": f"{filename}.pdf",
            "content_base64": content_base64,
        }

    except Exception as e:
        logger.error(f"导出PDF失败: {e}")
        return {"success": False, "error": str(e)}


@tool
async def export_to_docx(
    content: str,
    filename: str = "",
    title: str = "",
    format_type: Literal["markdown", "text"] = "markdown",
) -> dict[str, Any]:
    """将 Markdown 或纯文本生成可继续编辑的 Word 文档。"""
    try:
        import base64

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()
        normal = document.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)

        if title:
            heading = document.add_heading(_clean_markdown_inline(title), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        lines = content.splitlines()
        index = 0
        while index < len(lines):
            raw_line = lines[index].rstrip()
            line = raw_line.strip()
            if not line:
                index += 1
                continue

            if (
                format_type == "markdown"
                and "|" in line
                and index + 1 < len(lines)
                and re.match(r"^\|?\s*:?-{3,}", lines[index + 1].strip())
            ):
                headers = [cell.strip() for cell in line.strip("|").split("|")]
                rows: list[list[str]] = []
                index += 2
                while index < len(lines) and "|" in lines[index]:
                    rows.append(
                        [
                            _clean_markdown_inline(cell)
                            for cell in lines[index].strip().strip("|").split("|")
                        ]
                    )
                    index += 1
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for cell, value in zip(table.rows[0].cells, headers, strict=False):
                    cell.text = _clean_markdown_inline(value)
                for row in rows:
                    cells = table.add_row().cells
                    for cell_index, value in enumerate(row[: len(cells)]):
                        cells[cell_index].text = value
                continue

            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if format_type == "markdown" and heading_match:
                level = min(len(heading_match.group(1)), 4)
                document.add_heading(
                    _clean_markdown_inline(heading_match.group(2)), level=level
                )
            elif re.match(r"^[-*+]\s+", line):
                document.add_paragraph(
                    _clean_markdown_inline(re.sub(r"^[-*+]\s+", "", line)),
                    style="List Bullet",
                )
            elif re.match(r"^\d+[.)]\s+", line):
                document.add_paragraph(
                    _clean_markdown_inline(re.sub(r"^\d+[.)]\s+", "", line)),
                    style="List Number",
                )
            else:
                document.add_paragraph(_clean_markdown_inline(line.lstrip("> ")))
            index += 1

        output = io.BytesIO()
        document.save(output)
        content_base64 = base64.b64encode(output.getvalue()).decode()
        if not filename:
            filename = f"document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        return {
            "success": True,
            "filename": f"{filename}.docx",
            "content_base64": content_base64,
        }
    except Exception as e:  # broad-except: intentional
        logger.error("导出Word失败: %s", e)
        return {"success": False, "error": str(e)}
