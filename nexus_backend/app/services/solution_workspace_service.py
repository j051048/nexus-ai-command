"""Evidence-grounded solution generation and deterministic quality checks."""

from __future__ import annotations

import io
import json
import re
from datetime import UTC, datetime
from typing import Any
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from app.agent.artifact_contract import ArtifactSpec, ArtifactType
from app.agent.scientific_writing_skills import (
    build_writing_skill_prompt,
    enrich_artifact_spec,
)
from app.services.llm_gateway import llm_gateway
from app.services.scientific_instrument_domain import build_instrument_context
from app.services.solution_commercial_service import (
    enrich_workspace_commercials,
    extract_requirement_candidates,
)
from app.services.solution_quality_eval_service import evaluate_solution

WORKSPACE_SCHEMA_VERSION = "solution-workspace.v1"
STAGES = ["brief", "requirements", "configuration", "draft", "review", "delivery"]


def build_initial_workspace(brief: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_version": WORKSPACE_SCHEMA_VERSION,
        "active_stage": "brief",
        "brief": brief,
        "requirements": [],
        "packages": [],
        "sections": [],
        "review_gates": [
            {"id": "budget", "label": "预算范围已核对", "passed": False},
            {"id": "evidence", "label": "关键参数有企业资料依据", "passed": False},
            {"id": "claims", "label": "外部承诺已由负责人确认", "passed": False},
        ],
        "artifacts": [],
        "generation": {},
        "extension_data": {"output_connectors": [], "template_id": None},
    }


def apply_template_structure(
    workspace: dict[str, Any], template_structure: dict[str, Any]
) -> dict[str, Any]:
    """Reuse approved structure without carrying customer-specific approvals."""
    requirements = [
        {**item, "status": "open"}
        for item in template_structure.get("requirements") or []
        if isinstance(item, dict)
    ]
    sections = [
        {**item, "status": "draft"}
        for item in template_structure.get("sections") or []
        if isinstance(item, dict)
    ]
    review_gates = [
        {**item, "passed": False}
        for item in template_structure.get("review_gates") or []
        if isinstance(item, dict)
    ]
    templated = {
        **workspace,
        "requirements": requirements,
        "packages": [
            item
            for item in template_structure.get("packages") or []
            if isinstance(item, dict)
        ],
        "sections": sections,
        "review_gates": review_gates or workspace["review_gates"],
    }
    templated["quality"] = validate_workspace(templated)
    return templated


def _parse_json_object(content: str) -> dict[str, Any] | None:
    cleaned = content.strip()
    fenced = re.search(r"```(?:json)?\s*(\{.*\})\s*```", cleaned, re.DOTALL)
    if fenced:
        cleaned = fenced.group(1)
    else:
        start, end = cleaned.find("{"), cleaned.rfind("}")
        if start >= 0 and end > start:
            cleaned = cleaned[start : end + 1]
    try:
        parsed = json.loads(cleaned)
        return parsed if isinstance(parsed, dict) else None
    except (json.JSONDecodeError, TypeError):
        return None


def _fallback_draft(
    brief: dict[str, Any], products: list[dict[str, Any]]
) -> dict[str, Any]:
    budget_min = brief.get("budget_min")
    budget_max = brief.get("budget_max")
    budget_text = "预算待确认"
    if budget_min is not None or budget_max is not None:
        budget_text = f"预算范围：{budget_min or 0:,.0f} - {budget_max or 0:,.0f} 元"
    product_names = [
        item.get("model_code") or item.get("product_name") for item in products[:6]
    ]
    product_names = [name for name in product_names if name]
    requirements = [
        {
            "id": "req-1",
            "title": brief.get("application_scenario") or "明确应用目标与样品类型",
            "priority": "must",
            "status": "open",
            "evidence_ref": None,
        },
        {
            "id": "req-2",
            "title": budget_text,
            "priority": "must",
            "status": "open",
            "evidence_ref": None,
        },
        {
            "id": "req-3",
            "title": "确认安装环境、交付周期与验收标准",
            "priority": "should",
            "status": "open",
            "evidence_ref": None,
        },
    ]
    packages = []
    for key, name, positioning in [
        ("essential", "基础方案", "满足核心检测与合规要求"),
        ("recommended", "推荐方案", "兼顾性能、扩展性与全生命周期成本"),
        ("advanced", "进阶方案", "面向高通量、自动化与未来扩展"),
    ]:
        packages.append(
            {
                "id": key,
                "name": name,
                "positioning": positioning,
                "product_models": product_names[:2],
                "components": ["主机配置待选型", "安装培训", "基础应用支持"],
                "rationale": "需结合企业知识库中的产品参数进一步核验。",
                "tradeoffs": ["具体参数与价格待人工确认"],
            }
        )
    sections = [
        {
            "id": "summary",
            "title": "方案摘要",
            "content": f"面向{brief.get('customer_name') or '目标客户'}的{brief.get('application_scenario') or '科学仪器应用'}需求，提供分层配置建议。",
            "evidence_refs": [],
            "status": "draft",
        },
        {
            "id": "needs",
            "title": "需求理解",
            "content": "当前信息不足的项目已标记为待核验，建议在正式外发前完成技术澄清。",
            "evidence_refs": [],
            "status": "draft",
        },
        {
            "id": "delivery",
            "title": "交付与服务",
            "content": "交付周期、安装条件、培训范围与验收口径均需由业务负责人确认。",
            "evidence_refs": [],
            "status": "draft",
        },
    ]
    return {"requirements": requirements, "packages": packages, "sections": sections}


def validate_workspace(workspace: dict[str, Any]) -> dict[str, Any]:
    requirements = workspace.get("requirements") or []
    packages = workspace.get("packages") or []
    sections = workspace.get("sections") or []
    review_gates = workspace.get("review_gates") or []
    evidence_count = sum(len(item.get("evidence_refs") or []) for item in sections)
    open_claims = sum(
        1
        for item in requirements
        if item.get("priority") == "must"
        and (
            item.get("status") != "verified"
            or not str(item.get("evidence_ref") or "").strip()
        )
    )
    commercial = (workspace.get("extension_data") or {}).get("commercial_validation")
    commercial_valid = not isinstance(commercial, dict) or bool(commercial.get("valid"))
    checks = {
        "has_brief": bool((workspace.get("brief") or {}).get("title")),
        "has_requirements": bool(requirements),
        "has_three_packages": len(packages) >= 3,
        "has_draft": bool(sections),
        "sections_approved": bool(sections)
        and all(item.get("status") == "approved" for item in sections),
        "has_evidence": evidence_count > 0,
        "must_requirements_verified": open_claims == 0,
        "review_gates_passed": bool(review_gates)
        and all(bool(item.get("passed")) for item in review_gates),
        "commercial_configuration_valid": commercial_valid,
    }
    deterministic_quality = evaluate_solution(workspace, stage="external")
    checks["deterministic_quality_ready"] = deterministic_quality["ready"]
    return {
        "checks": checks,
        "ready_for_external_use": all(checks.values()),
        "evidence_count": evidence_count,
        "open_claims": open_claims,
        "completion": round(sum(checks.values()) / len(checks) * 100),
        "deterministic_evaluation": deterministic_quality,
    }


async def extract_requirements(
    *,
    documents: list[dict[str, Any]],
    brief: dict[str, Any],
    user_id: str,
    organization_id: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Extract a traceable requirement matrix, with deterministic fallback."""
    fallback = extract_requirement_candidates(documents)
    document_payload = []
    for document in documents[:12]:
        extracted = document.get("extracted_data") or {}
        if isinstance(extracted, dict):
            content = extracted.get("full_text_context") or extracted.get("summary")
        else:
            content = str(extracted)
        document_payload.append(
            {
                "id": str(document.get("id") or ""),
                "name": document.get("name"),
                "doc_type": document.get("doc_type"),
                "source_version": document.get("source_version"),
                "valid_until": document.get("valid_until"),
                "content": str(content or "")[:10000],
            }
        )
    if not document_payload:
        return fallback, {"degraded": True, "reason": "no_documents"}
    response = await llm_gateway.chat(
        scene_code="solution_requirement_extraction",
        agent_code="scientific_requirement_analyst",
        user_id=user_id,
        org_id=organization_id,
        system_prompt=(
            "你是科学仪器售前需求分析师。只提取资料中明确存在的需求，不补造参数。"
            "输出 JSON 对象 {requirements: []}。每项必须包含 title、priority、"
            "source_document_id、source_name、source_excerpt；priority 只能是 must、"
            "should、optional。否决项、必须项、强制项归类为 must。"
        ),
        messages=[
            {
                "role": "user",
                "content": json.dumps(
                    {"brief": brief, "documents": document_payload},
                    ensure_ascii=False,
                    default=str,
                ),
            }
        ],
        temperature=0,
        max_tokens=2600,
    )
    parsed = (
        _parse_json_object(response.content)
        if response.finish_reason != "error"
        else None
    )
    rows = parsed.get("requirements") if isinstance(parsed, dict) else None
    if not isinstance(rows, list):
        return fallback, {
            "degraded": True,
            "model": response.model_code,
            "usage": response.usage,
        }
    document_names = {
        str(item.get("id")): item.get("name") for item in document_payload
    }
    normalized: list[dict[str, Any]] = []
    for index, row in enumerate(rows[:100]):
        if not isinstance(row, dict) or not str(row.get("title") or "").strip():
            continue
        source_id = str(row.get("source_document_id") or "")
        normalized.append(
            {
                "id": f"req-doc-{index + 1}",
                "title": str(row["title"]).strip()[:500],
                "priority": (
                    row.get("priority")
                    if row.get("priority") in {"must", "should", "optional"}
                    else "should"
                ),
                "status": "open",
                "evidence_ref": source_id or None,
                "source_document_id": source_id or None,
                "source_name": row.get("source_name") or document_names.get(source_id),
                "source_excerpt": str(row.get("source_excerpt") or "")[:500],
            }
        )
    return normalized or fallback, {
        "degraded": not bool(normalized),
        "model": response.model_code,
        "usage": response.usage,
    }


async def generate_solution(
    *,
    brief: dict[str, Any],
    current_workspace: dict[str, Any],
    products: list[dict[str, Any]],
    knowledge_context: str,
    template: dict[str, Any] | None,
    user_id: str,
    organization_id: str,
) -> tuple[dict[str, Any], dict[str, Any]]:
    source_context = knowledge_context[:12000]
    product_context = json.dumps(products[:12], ensure_ascii=False, default=str)[:8000]
    template_context = json.dumps(template or {}, ensure_ascii=False, default=str)[
        :5000
    ]
    domain_context = build_instrument_context(
        brief.get("instrument_line_code"),
        application_field=brief.get("application_scenario"),
    )
    artifact_spec = enrich_artifact_spec(
        current_workspace.get("artifact_spec")
        or ArtifactSpec(
            artifact_type=ArtifactType.CUSTOMER_SOLUTION,
            external_delivery=True,
            strict_quality=True,
            instrument_line=brief.get("instrument_line_code"),
            industry=brief.get("industry"),
            region=brief.get("region"),
            requested_formats=["docx", "pdf", "xlsx"],
        )
    )
    writing_contract = build_writing_skill_prompt(artifact_spec)
    system_prompt = """你是科学仪器企业的资深售前方案架构师，覆盖光谱、色谱、质谱、能谱及电子测试仪器。
只允许依据输入的客户事实、产品目录和企业知识资料写作。不得虚构参数、资质、交付周期、价格或客户案例；证据不足时明确写“待核验”。
输出严格 JSON 对象，字段仅包含 requirements、packages、sections。packages 必须提供基础、推荐、进阶三档；sections 中每段必须有 evidence_refs 数组。
外发、报价和承诺均由人工确认，本次只生成可编辑草稿。"""
    system_prompt += "\n\n" + writing_contract
    prompt = {
        "brief": brief,
        "existing_workspace": current_workspace,
        "product_catalog": product_context,
        "enterprise_knowledge": source_context,
        "approved_template": template_context,
        "instrument_domain_playbook": domain_context,
        "artifact_spec": artifact_spec.model_dump(mode="json"),
        "output_schema": {
            "requirements": [
                {
                    "id": "req-1",
                    "title": "",
                    "priority": "must|should|optional",
                    "status": "open",
                    "evidence_ref": "资料名或null",
                }
            ],
            "packages": [
                {
                    "id": "essential|recommended|advanced",
                    "name": "",
                    "positioning": "",
                    "product_models": [],
                    "components": [],
                    "rationale": "",
                    "tradeoffs": [],
                }
            ],
            "sections": [
                {
                    "id": "",
                    "title": "",
                    "content": "",
                    "evidence_refs": [],
                    "status": "draft",
                }
            ],
        },
    }
    response = await llm_gateway.chat(
        scene_code="solution_design",
        agent_code="scientific_solution_architect",
        user_id=user_id,
        org_id=organization_id,
        system_prompt=system_prompt,
        messages=[
            {
                "role": "user",
                "content": json.dumps(prompt, ensure_ascii=False, default=str),
            }
        ],
        temperature=0.2,
        max_tokens=3200,
    )
    generated = (
        _parse_json_object(response.content)
        if response.finish_reason != "error"
        else None
    )
    degraded = generated is None
    generated = generated or _fallback_draft(brief, products)
    workspace = enrich_workspace_commercials(
        {
            **current_workspace,
            **generated,
            "schema_version": WORKSPACE_SCHEMA_VERSION,
            "artifact_spec": artifact_spec.model_dump(mode="json"),
            "active_stage": "review",
            "generation": {
                "generated_at": datetime.now(UTC).isoformat(),
                "model": response.model_code,
                "usage": response.usage,
                "degraded": degraded,
                "knowledge_context_available": bool(knowledge_context),
            },
        },
        products,
    )
    quality = evaluate_solution(workspace, stage="draft")
    repair_count = 0
    evidence_packet = (workspace.get("extension_data") or {}).get("evidence_packet", {})
    while (
        repair_count < artifact_spec.max_repair_cycles
        and not quality.get("artifact_quality", {}).get("ready", False)
        and evidence_packet.get("sufficient", False)
        and quality.get("repairable_findings")
    ):
        repair_count += 1
        repair_response = await llm_gateway.chat(
            scene_code="solution_section_repair",
            agent_code="scientific_solution_editor",
            user_id=user_id,
            org_id=organization_id,
            system_prompt=(
                writing_contract
                + "\nRepair only the sections. Return strict JSON: {sections: [...]}。"
            ),
            messages=[
                {
                    "role": "user",
                    "content": json.dumps(
                        {
                            "sections": workspace.get("sections") or [],
                            "findings": quality.get("repairable_findings") or [],
                            "enterprise_knowledge": source_context,
                        },
                        ensure_ascii=False,
                        default=str,
                    ),
                }
            ],
            temperature=0.1,
            max_tokens=2600,
        )
        repaired = (
            _parse_json_object(repair_response.content)
            if repair_response.finish_reason != "error"
            else None
        )
        if not repaired or not isinstance(repaired.get("sections"), list):
            break
        workspace["sections"] = repaired["sections"]
        quality = evaluate_solution(workspace, stage="draft")

    workspace["generation"]["repair_count"] = repair_count
    workspace["generation"]["artifact_quality_score"] = quality.get("score", 0)
    workspace["quality"] = {
        **validate_workspace(workspace),
        "deterministic_evaluation": quality,
    }
    return workspace, workspace["generation"]


def workspace_markdown(project: dict[str, Any]) -> str:
    workspace = project.get("workspace") or {}
    lines = [f"# {project.get('title') or '客户解决方案'}", ""]
    brief = workspace.get("brief") or {}
    lines.extend(
        [
            f"客户：{brief.get('customer_name') or '待确认'}",
            f"行业/地区：{brief.get('industry') or '待确认'} / {brief.get('region') or '待确认'}",
            "",
        ]
    )
    requirements = workspace.get("requirements") or []
    if requirements:
        lines.extend(["## 需求与验收矩阵", ""])
        for item in requirements:
            lines.append(
                "- [{status}] {title} | 优先级: {priority} | 证据: {evidence}".format(
                    status=item.get("status") or "open",
                    title=item.get("title") or "未命名需求",
                    priority=item.get("priority") or "should",
                    evidence=item.get("evidence_ref") or "待核验",
                )
            )
        lines.append("")
    for section in workspace.get("sections") or []:
        lines.extend(
            [
                f"## {section.get('title') or '未命名章节'}",
                section.get("content") or "",
                "",
            ]
        )
        refs = section.get("evidence_refs") or []
        if refs:
            lines.extend([f"证据：{', '.join(map(str, refs))}", ""])
    lines.extend(["## 配置建议", ""])
    for package in workspace.get("packages") or []:
        lines.extend(
            [f"### {package.get('name')}", package.get("positioning") or "", ""]
        )
        for component in package.get("components") or []:
            lines.append(f"- {component}")
        lines.append("")
    evidence_catalog = (workspace.get("extension_data") or {}).get(
        "evidence_catalog", []
    )
    if evidence_catalog:
        lines.extend(["## 证据附录", ""])
        for item in evidence_catalog[:20]:
            citation = (
                f"EVID:{item.get('document_id')}:{item.get('chunk_id')}"
                if item.get("document_id") and item.get("chunk_id")
                else item.get("document_id") or "待核验"
            )
            version = item.get("source_version") or "未标注版本"
            lines.append(
                f"- [{citation}] {item.get('title') or item.get('source') or '企业资料'} ({version})"
            )
        lines.append("")
    lines.append("> 本文档为 AI 辅助草稿，参数、价格、交期与外部承诺须经人工审核。")
    return "\n".join(lines)


def export_docx(project: dict[str, Any], brand: dict[str, Any] | None = None) -> bytes:
    markdown = workspace_markdown(project)
    document = Document()
    brand = brand or {}
    title = document.add_heading(project.get("title") or "客户解决方案", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = document.add_paragraph(
        str(brand.get("company_name") or brand.get("name") or "企业解决方案")
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph(
        f"项目编号：{project.get('project_code') or '-'}    "
        f"版本：v{project.get('current_version') or 0}"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_page_break()
    for line in markdown.splitlines():
        if line.startswith("# "):
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
        elif line.startswith("> "):
            document.add_paragraph(line[2:])
        else:
            document.add_paragraph(line)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_xlsx(project: dict[str, Any]) -> bytes:
    """Export requirement and configuration matrices for internal review."""
    workspace = project.get("workspace") or {}
    workbook = Workbook()
    requirements_sheet = workbook.active
    requirements_sheet.title = "Requirement Matrix"
    requirement_headers = [
        "Priority",
        "Requirement",
        "Status",
        "Evidence",
        "Source",
        "Excerpt",
    ]
    requirements_sheet.append(requirement_headers)
    for requirement in workspace.get("requirements") or []:
        requirements_sheet.append(
            [
                requirement.get("priority"),
                requirement.get("title"),
                requirement.get("status"),
                requirement.get("evidence_ref"),
                requirement.get("source_name"),
                requirement.get("source_excerpt"),
            ]
        )

    configuration_sheet = workbook.create_sheet("Configuration")
    configuration_headers = [
        "Package",
        "Positioning",
        "Models",
        "Components",
        "Currency",
        "List Price",
        "Standard Cost",
        "Gross Margin %",
        "Lead Time Days",
        "Warnings",
    ]
    configuration_sheet.append(configuration_headers)
    for package in workspace.get("packages") or []:
        commercial = package.get("commercial") or {}
        configuration_sheet.append(
            [
                package.get("name"),
                package.get("positioning"),
                "\n".join(package.get("product_models") or []),
                "\n".join(package.get("components") or []),
                commercial.get("currency"),
                commercial.get("list_price"),
                commercial.get("standard_cost"),
                commercial.get("gross_margin_percent"),
                commercial.get("lead_time_days"),
                "\n".join(commercial.get("validation_warnings") or []),
            ]
        )

    evidence_sheet = workbook.create_sheet("Evidence")
    evidence_sheet.append(
        ["Document", "Title", "Type", "Version", "Valid Until", "Score", "Excerpt"]
    )
    evidence_catalog = (workspace.get("extension_data") or {}).get(
        "evidence_catalog"
    ) or []
    for evidence in evidence_catalog:
        evidence_sheet.append(
            [
                evidence.get("document_id"),
                evidence.get("title"),
                evidence.get("doc_type"),
                evidence.get("source_version"),
                evidence.get("valid_until"),
                evidence.get("score"),
                evidence.get("excerpt"),
            ]
        )

    for sheet in workbook.worksheets:
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="1F4B6E")
            cell.alignment = Alignment(vertical="center")
        sheet.freeze_panes = "A2"
        sheet.auto_filter.ref = sheet.dimensions
        for column in sheet.columns:
            letter = column[0].column_letter
            max_length = min(
                48,
                max(len(str(cell.value or "")) for cell in column) + 2,
            )
            sheet.column_dimensions[letter].width = max(12, max_length)
        for row in sheet.iter_rows(min_row=2):
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()


def export_pdf(project: dict[str, Any]) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = "STSong-Light"
    story = []
    for line in workspace_markdown(project).splitlines():
        if not line:
            story.append(Spacer(1, 3 * mm))
            continue
        if line.startswith("# "):
            story.append(Paragraph(escape(line[2:]), styles["Title"]))
        elif line.startswith("## "):
            story.append(Paragraph(escape(line[3:]), styles["Heading2"]))
        elif line.startswith("### "):
            story.append(Paragraph(escape(line[4:]), styles["Heading3"]))
        else:
            story.append(Paragraph(escape(line.lstrip("- >")), styles["BodyText"]))
    SimpleDocTemplate(
        buffer, pagesize=A4, rightMargin=18 * mm, leftMargin=18 * mm
    ).build(story)
    return buffer.getvalue()
