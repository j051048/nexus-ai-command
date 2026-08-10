"""Delivery-safety scanning for externally sent artifacts.

Covers PII leakage, internal markers, promise/commitment language and a
lightweight DOCX render sanity check.  Deterministic by design; the LLM
judge covers semantic safety, this module covers mechanical safety.
"""

from __future__ import annotations

import re
from typing import Any

DELIVERY_SCAN_VERSION = "artifact-delivery-scan.v1"

_PHONE_RE = re.compile(r"(?<!\d)1[3-9]\d{9}(?!\d)")
_ID_CARD_RE = re.compile(r"(?<!\d)\d{17}[\dXx](?!\d)")
_EMAIL_RE = re.compile(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}")
_INTERNAL_MARKERS = ("内部资料", "内部文件", "机密", "TODO", "占位", "示例数据", "XXX")
_PROMISE_PATTERNS = (
    re.compile(r"(保证|确保|承诺)[^。；;]{0,24}(交付|完成|上线|达标)"),
    re.compile(r"\d+\s*个?(工作日|自然日|天内)\s*(交付|完成|上线)"),
    re.compile(r"(绝对|百分之百|100%)"),
)
_COMPLIANCE_TERMS = ("投标", "招标", "监管", "政策", "合规")


def scan_delivery_safety(text: str) -> dict[str, Any]:
    """Scan an artifact for mechanical delivery risks."""
    text = str(text or "")
    findings: list[dict[str, Any]] = []
    details: dict[str, Any] = {}

    phones = _PHONE_RE.findall(text)
    if phones:
        findings.append(
            {
                "severity": "high",
                "code": "pii_phone",
                "message": f"检测到 {len(phones)} 处手机号，外发前需确认授权",
            }
        )
        details["phone_count"] = len(phones)

    id_cards = _ID_CARD_RE.findall(text)
    if id_cards:
        findings.append(
            {
                "severity": "high",
                "code": "pii_id_card",
                "message": f"检测到 {len(id_cards)} 处身份证号，外发前必须移除或脱敏",
            }
        )
        details["id_card_count"] = len(id_cards)

    emails = _EMAIL_RE.findall(text)
    if emails:
        findings.append(
            {
                "severity": "medium",
                "code": "pii_email",
                "message": f"检测到 {len(emails)} 处邮箱地址",
            }
        )
        details["email_count"] = len(emails)

    internal_hits = [marker for marker in _INTERNAL_MARKERS if marker in text]
    if internal_hits:
        findings.append(
            {
                "severity": "high",
                "code": "internal_marker",
                "message": f"外发文档包含内部标记：{', '.join(internal_hits)}",
            }
        )
        details["internal_markers"] = internal_hits

    promise_hits: list[str] = []
    for pattern in _PROMISE_PATTERNS:
        for match in pattern.findall(text):
            if isinstance(match, tuple):
                match = "".join(match)
            promise_hits.append(match)
    if promise_hits:
        findings.append(
            {
                "severity": "medium",
                "code": "unverified_promise",
                "message": f"检测到 {len(promise_hits)} 处承诺性表述，需人工确认边界",
                "details": {"samples": promise_hits[:5]},
            }
        )
        details["promise_count"] = len(promise_hits)

    compliance = [term for term in _COMPLIANCE_TERMS if term in text]
    if compliance:
        details["compliance_terms"] = compliance

    deductions: dict[str, float] = {"high": 30.0, "medium": 12.0}
    score = max(0.0, 100.0 - sum(deductions[f["severity"]] for f in findings))
    return {
        "evaluator_version": DELIVERY_SCAN_VERSION,
        "score": round(score, 2),
        "findings": findings,
        "details": details,
    }


def verify_docx_render(docx_bytes: bytes) -> dict[str, Any]:
    """Sanity-check a rendered DOCX: non-empty paragraphs and tables."""
    try:
        from docx import Document

        document = Document(docx_bytes)
        paragraphs = len(document.paragraphs)
        tables = len(document.tables)
        text_len = sum(len(p.text or "") for p in document.paragraphs)
        if paragraphs == 0 or text_len == 0:
            return {
                "ok": False,
                "error": "rendered docx is empty",
                "paragraphs": paragraphs,
                "tables": tables,
            }
        return {
            "ok": True,
            "paragraphs": paragraphs,
            "tables": tables,
            "character_count": text_len,
        }
    except Exception as exc:  # broad-except: intentional
        return {"ok": False, "error": str(exc)}
