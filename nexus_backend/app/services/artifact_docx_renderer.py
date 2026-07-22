"""Professional, deterministic renderers for reviewed Agent deliverables."""

from __future__ import annotations

import io
import re
from collections.abc import Iterable
from datetime import UTC, datetime
from html import escape
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill

from app.services.artifact_content_sanitizer import sanitize_artifact_content

BRAND_BLUE = "164E63"
ACCENT_BLUE = "0F766E"
LIGHT_BLUE = "E8F2F3"
PALE_BLUE = "F3F8F8"
TABLE_BORDER = "CBD5E1"
TEXT_DARK = "182230"
TEXT_MUTED = "667085"


def _cover_title(value: str) -> str:
    """Split long solution titles at a meaningful business suffix."""

    title = " ".join(value.split()).strip()
    if len(title) < 22:
        return title
    for suffix in (
        "升级换代整体解决方案",
        "综合解决方案",
        "整体解决方案",
        "技术解决方案",
        "解决方案",
    ):
        position = title.rfind(suffix)
        if position >= 8:
            return f"{title[:position].rstrip(' -—：:')}\n{title[position:]}"
    return title


def _set_cell_shading(cell: Any, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_cell_margins(
    cell: Any, *, top: int = 100, start: int = 120, bottom: int = 100, end: int = 120
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_table_geometry(table: Any, widths: list[int]) -> None:
    properties = table._tbl.tblPr
    width = properties.first_child_found_in("w:tblW")
    if width is None:
        width = OxmlElement("w:tblW")
        properties.append(width)
    width.set(qn("w:w"), str(sum(widths)))
    width.set(qn("w:type"), "dxa")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for column_width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(column_width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Cm(widths[index] / 1440 * 2.54)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_properties = cell._tc.get_or_add_tcPr()
            cell_width = cell_properties.first_child_found_in("w:tcW")
            if cell_width is None:
                cell_width = OxmlElement("w:tcW")
                cell_properties.append(cell_width)
            cell_width.set(qn("w:w"), str(widths[index]))
            cell_width.set(qn("w:type"), "dxa")
            _set_cell_margins(cell)


def _set_table_borders(table: Any, color: str = TABLE_BORDER) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def _column_widths(rows: list[list[str]], column_count: int) -> list[int]:
    if column_count <= 1:
        return [9360]
    weights: list[float] = []
    for column in range(column_count):
        longest = max(
            (len(row[column]) if column < len(row) else 0 for row in rows),
            default=1,
        )
        weights.append(max(0.9, min(3.4, longest / 12)))
    total = sum(weights)
    widths = [int(9360 * weight / total) for weight in weights]
    widths[-1] += 9360 - sum(widths)
    return widths


def _set_paragraph_shading(paragraph: Any, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _set_left_border(paragraph: Any, color: str, size: int = 18) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = properties.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        properties.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)


def _add_page_number(paragraph: Any) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    run.font.size = Pt(8)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    paragraph.add_run(" 页").font.size = Pt(8)


def _configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for name, size, color in (
        ("Title", 28, BRAND_BLUE),
        ("Heading 1", 16, BRAND_BLUE),
        ("Heading 2", 13, ACCENT_BLUE),
        ("Heading 3", 11.5, TEXT_DARK),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.keep_with_next = True

    if "Source Note" not in styles:
        source_style = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
        source_style.base_style = styles["Normal"]
        source_style.font.size = Pt(8.5)
        source_style.font.color.rgb = RGBColor.from_string(TEXT_MUTED)

    if "Executive Summary" not in styles:
        summary_style = styles.add_style("Executive Summary", WD_STYLE_TYPE.PARAGRAPH)
        summary_style.base_style = styles["Normal"]
        summary_style.font.size = Pt(11)
        summary_style.font.color.rgb = RGBColor.from_string(TEXT_DARK)
        summary_style.paragraph_format.left_indent = Cm(0.45)
        summary_style.paragraph_format.right_indent = Cm(0.45)
        summary_style.paragraph_format.space_before = Pt(4)
        summary_style.paragraph_format.space_after = Pt(10)


def _split_markdown_table(
    lines: list[str], start: int
) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines) or "|" not in lines[start]:
        return None
    separator = lines[start + 1].strip()
    if not re.match(r"^\|?\s*:?-{3,}", separator):
        return None
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        row = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        if index != start + 1:
            rows.append(row)
        index += 1
    return rows, index


def _add_rich_text(paragraph: Any, value: str) -> None:
    token_re = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|\[来源\s+\d+\])")
    position = 0
    for match in token_re.finditer(value):
        if match.start() > position:
            paragraph.add_run(value[position : match.start()])
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
            run.font.color.rgb = RGBColor.from_string(BRAND_BLUE)
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(token)
            run.bold = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor.from_string(ACCENT_BLUE)
        position = match.end()
    if position < len(value):
        paragraph.add_run(value[position:])


def _mark_table_header(row: Any) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def _render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    index = 0
    current_heading = ""
    while index < len(lines):
        line = lines[index].rstrip()
        table_data = _split_markdown_table(lines, index)
        if table_data:
            rows, next_index = table_data
            if rows:
                width = max(len(row) for row in rows)
                table = document.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                _set_table_geometry(table, _column_widths(rows, width))
                _set_table_borders(table)
                _mark_table_header(table.rows[0])
                for row_index, row in enumerate(rows):
                    for column_index in range(width):
                        value = row[column_index] if column_index < len(row) else ""
                        cell = table.cell(row_index, column_index)
                        cell.text = ""
                        cell_paragraph = cell.paragraphs[0]
                        cell_paragraph.paragraph_format.space_after = Pt(0)
                        cell_paragraph.paragraph_format.line_spacing = 1.15
                        _add_rich_text(cell_paragraph, value)
                        if row_index == 0:
                            _set_cell_shading(cell, BRAND_BLUE)
                            for run in cell.paragraphs[0].runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                        elif row_index % 2 == 0:
                            _set_cell_shading(cell, PALE_BLUE)
                document.add_paragraph()
            index = next_index
            continue
        stripped = line.strip()
        if not stripped:
            index += 1
            continue
        if stripped.startswith("### "):
            current_heading = stripped[4:].strip()
            heading = document.add_heading(current_heading, level=3)
            _set_left_border(heading, TABLE_BORDER, 8)
        elif stripped.startswith("## "):
            current_heading = stripped[3:].strip()
            heading = document.add_heading(current_heading, level=2)
            _set_left_border(heading, ACCENT_BLUE, 16)
        elif stripped.startswith("# "):
            current_heading = stripped[2:].strip()
            document.add_heading(current_heading, level=1)
        elif re.match(r"^[-*]\s+", stripped):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_rich_text(paragraph, re.sub(r"^[-*]\s+", "", stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _add_rich_text(paragraph, re.sub(r"^\d+[.)]\s+", "", stripped))
        elif stripped.startswith("> "):
            paragraph = document.add_paragraph(style="Executive Summary")
            _add_rich_text(paragraph, stripped[2:])
            paragraph.paragraph_format.left_indent = Cm(0.6)
            paragraph.paragraph_format.right_indent = Cm(0.6)
            _set_paragraph_shading(paragraph, PALE_BLUE)
            _set_left_border(paragraph, ACCENT_BLUE, 18)
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        else:
            is_summary = current_heading == "执行摘要"
            paragraph = document.add_paragraph(
                style="Executive Summary" if is_summary else None
            )
            _add_rich_text(paragraph, stripped)
            if is_summary:
                _set_paragraph_shading(paragraph, PALE_BLUE)
                _set_left_border(paragraph, ACCENT_BLUE, 18)
                current_heading = ""
        index += 1


def render_artifact_docx(
    artifact: dict[str, Any],
    evidence_packet: dict[str, Any] | None = None,
    brand: dict[str, Any] | None = None,
) -> bytes:
    """Render a customer-ready Word document from the canonical artifact version."""

    brand = brand or {}
    sanitized = sanitize_artifact_content(
        str(artifact.get("content_markdown") or ""), evidence_packet
    )
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(1.9)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.0)
    section.different_first_page_header_footer = True
    _configure_styles(document)

    company = str(brand.get("company_name") or brand.get("name") or "Nexus AI")
    title = str(artifact.get("title") or "AI 成果")
    cover_label = str(artifact.get("artifact_label") or "科学仪器专业成果")
    paragraph = document.add_paragraph(company.upper())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    paragraph.paragraph_format.space_after = Pt(20)
    for run in paragraph.runs:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(ACCENT_BLUE)
    title_paragraph = document.add_paragraph(_cover_title(title), style="Title")
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(12)
    title_size = 22 if len(title) > 34 else 24 if len(title) > 24 else 28
    for run in title_paragraph.runs:
        run.font.size = Pt(title_size)
    subtitle = document.add_paragraph(cover_label)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
    document.add_paragraph()
    meta_table = document.add_table(rows=4, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    _set_table_geometry(meta_table, [2480, 6880])
    _set_table_borders(meta_table)
    meta_rows = (
        ("成果编号", str(artifact.get("artifact_code") or "-")),
        ("版本", f"v{artifact.get('version_number') or 1}"),
        (
            "状态",
            (
                "经人工确认"
                if artifact.get("approval_status") == "approved"
                else "审核草稿"
            ),
        ),
        ("生成日期", datetime.now(UTC).strftime("%Y-%m-%d")),
    )
    for row_index, (label, value) in enumerate(meta_rows):
        meta_table.cell(row_index, 0).text = label
        meta_table.cell(row_index, 1).text = value
        _set_cell_shading(meta_table.cell(row_index, 0), LIGHT_BLUE)
        meta_table.cell(row_index, 0).paragraphs[0].runs[0].font.bold = True

    quality = artifact.get("quality_snapshot") or {}
    quality_score = float(artifact.get("quality_score") or quality.get("score") or 0)
    evidence_count = len((evidence_packet or {}).get("records") or [])
    metrics_table = document.add_table(rows=1, cols=3)
    metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics_table.autofit = False
    _set_table_geometry(metrics_table, [3120, 3120, 3120])
    _set_table_borders(metrics_table)
    metrics = (
        (str(evidence_count), "条企业证据"),
        (str(round(quality_score)), "质量评分"),
        (f"v{artifact.get('version_number') or 1}", "交付版本"),
    )
    for index, (value, label) in enumerate(metrics):
        cell = metrics_table.cell(0, index)
        cell.text = ""
        metric_paragraph = cell.paragraphs[0]
        metric_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        value_run = metric_paragraph.add_run(value)
        value_run.bold = True
        value_run.font.size = Pt(15)
        value_run.font.color.rgb = RGBColor.from_string(BRAND_BLUE)
        label_run = metric_paragraph.add_run(f"\n{label}")
        label_run.font.size = Pt(8.5)
        label_run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        _set_cell_shading(cell, PALE_BLUE)
    document.add_paragraph()
    notice = document.add_paragraph(
        "本成果由 AI 基于企业已授权资料辅助生成。参数、价格、交期、案例和对外承诺以最终人工审核版本为准。"
    )
    notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in notice.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
    _set_paragraph_shading(notice, PALE_BLUE)
    _set_left_border(notice, ACCENT_BLUE, 16)
    document.add_page_break()

    body_markdown = re.sub(
        r"^\s*#\s+.*?(?:\n+|$)", "", sanitized.content, count=1
    ).strip()
    headings = [
        re.sub(r"^#{1,3}\s+", "", line).strip()
        for line in body_markdown.splitlines()
        if re.match(r"^##{1,2}\s+", line)
    ]
    if headings:
        document.add_heading("目录", level=1)
        for position, heading in enumerate(headings[:24], 1):
            document.add_paragraph(f"{position:02d}  {heading}")
        document.add_page_break()

    _render_markdown(document, body_markdown)
    if sanitized.source_notes:
        document.add_page_break()
        document.add_heading("资料来源", level=1)
        for note in sanitized.source_notes:
            document.add_paragraph(
                f"[{note['number']}] {note['title']} · 版本 {note['version']}",
                style="Source Note",
            )

    for doc_section in document.sections:
        doc_section.different_first_page_header_footer = True
        header = doc_section.header.paragraphs[0]
        header.text = f"{company}  |  {title}"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string(TEXT_MUTED)
        _add_page_number(doc_section.footer.paragraphs[0])
        doc_section.first_page_header.paragraphs[0].text = ""
        doc_section.first_page_footer.paragraphs[0].text = ""

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _pdf_rows(markdown: str) -> Iterable[tuple[str, str]]:
    for line in markdown.splitlines():
        value = line.strip()
        if not value or re.match(r"^\|?\s*:?-{3,}", value):
            continue
        if value.startswith("### "):
            yield "Heading3", value[4:]
        elif value.startswith("## "):
            yield "Heading2", value[3:]
        elif value.startswith("# "):
            yield "Heading1", value[2:]
        else:
            yield "BodyText", value.lstrip("-*>").strip()


def render_artifact_pdf(
    artifact: dict[str, Any],
    evidence_packet: dict[str, Any] | None = None,
    brand: dict[str, Any] | None = None,
) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    sanitized = sanitize_artifact_content(
        str(artifact.get("content_markdown") or ""), evidence_packet
    )
    buffer = io.BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style in styles.byName.values():
        style.fontName = "STSong-Light"
    styles.add(
        ParagraphStyle(
            name="NexusNotice",
            parent=styles["BodyText"],
            fontName="STSong-Light",
            fontSize=8.5,
            leading=14,
            textColor=colors.HexColor("#667085"),
        )
    )
    title = str(artifact.get("title") or "AI 成果")
    company = str((brand or {}).get("company_name") or "Nexus AI")
    story = [
        Spacer(1, 38 * mm),
        Paragraph(escape(company), styles["Heading3"]),
        Spacer(1, 8 * mm),
        Paragraph(escape(title), styles["Title"]),
        Spacer(1, 8 * mm),
        Paragraph("科学仪器专业成果", styles["Heading2"]),
        Spacer(1, 24 * mm),
        Paragraph(
            "本成果由 AI 基于企业已授权资料辅助生成，外发前须完成事实与承诺复核。",
            styles["NexusNotice"],
        ),
        PageBreak(),
    ]
    for style_name, value in _pdf_rows(sanitized.content):
        story.append(Paragraph(escape(value), styles[style_name]))
        story.append(Spacer(1, 2.5 * mm))
    if sanitized.source_notes:
        story.extend([PageBreak(), Paragraph("资料来源", styles["Heading1"])])
        for note in sanitized.source_notes:
            story.append(
                Paragraph(
                    escape(
                        f"[{note['number']}] {note['title']} · 版本 {note['version']}"
                    ),
                    styles["NexusNotice"],
                )
            )
    document = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=18 * mm,
        leftMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
        author=company,
    )
    document.build(story)
    return buffer.getvalue()


def render_artifact_xlsx(
    artifact: dict[str, Any], evidence_packet: dict[str, Any] | None = None
) -> bytes:
    sanitized = sanitize_artifact_content(
        str(artifact.get("content_markdown") or ""), evidence_packet
    )
    workbook = Workbook()
    summary = workbook.active
    summary.title = "成果摘要"
    summary.append(["成果名称", artifact.get("title")])
    summary.append(["版本", artifact.get("version_number") or 1])
    summary.append(["质量评分", artifact.get("quality_score")])
    summary.append(["审批状态", artifact.get("approval_status")])
    summary.append([])
    summary.append(["正文"])
    for line in sanitized.content.splitlines():
        if line.strip():
            summary.append([line.strip()])

    evidence = workbook.create_sheet("证据清单")
    evidence.append(["序号", "资料", "版本", "文档 ID", "片段 ID"])
    for note in sanitized.source_notes:
        evidence.append(
            [
                note["number"],
                note["title"],
                note["version"],
                note["document_id"],
                note["chunk_id"],
            ]
        )
    verification = workbook.create_sheet("待核验项")
    verification.append(["序号", "待核验内容", "处理状态"])
    for index, line in enumerate(
        [item for item in sanitized.content.splitlines() if "待核验" in item], 1
    ):
        verification.append([index, line.strip(), "待确认"])

    for sheet in workbook.worksheets:
        sheet.freeze_panes = "A2"
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=BRAND_BLUE)
        for row in sheet.iter_rows():
            for cell in row:
                cell.alignment = Alignment(vertical="top", wrap_text=True)
        for column in sheet.columns:
            letter = column[0].column_letter
            sheet.column_dimensions[letter].width = min(
                72, max(14, max(len(str(cell.value or "")) for cell in column) + 2)
            )
    buffer = io.BytesIO()
    workbook.save(buffer)
    return buffer.getvalue()
