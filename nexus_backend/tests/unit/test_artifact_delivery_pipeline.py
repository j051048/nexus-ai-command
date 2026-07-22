import io
import json
from pathlib import Path
from types import SimpleNamespace

import pytest
from docx import Document

from app.agent.artifact_contract import (
    ArtifactAudience,
    ArtifactSpec,
    ArtifactType,
    infer_artifact_spec,
)
from app.agent.scientific_writing_skills import enrich_artifact_spec
from app.services.agent_evidence_service import EvidencePacket, EvidenceRecord
from app.services.artifact_content_sanitizer import (
    contains_internal_trace_markers,
    sanitize_artifact_content,
)
from app.services.artifact_docx_renderer import (
    render_artifact_docx,
    render_artifact_pdf,
    render_artifact_xlsx,
)
from app.services.artifact_evidence_compiler import compile_artifact_evidence
from app.services.artifact_generation_service import generate_artifact
from app.services.artifact_quality_service import evaluate_text_artifact


def _packet() -> dict:
    return {
        "records": [
            {
                "document_id": "doc-1",
                "chunk_id": "chunk-1",
                "title": "FD-F 产品手册",
                "source_version": "2026.2",
                "excerpt": "产品参数已经过企业审核。",
            }
        ],
        "coverage": 1.0,
        "sufficient": True,
        "missing_topics": [],
    }


def test_sanitizer_removes_trace_deduplicates_and_humanizes_sources():
    repeated = "这是一段需要去重的企业资料说明，内容足够长以便识别重复段落。"
    raw = (
        "[企业资料检索结果]\n\n"
        "tool_name: loadknowledge\n\n"
        f"{repeated}\n\n{repeated}\n\n"
        "核心参数已核验。[EVID:doc-1:chunk-1]"
    )

    result = sanitize_artifact_content(raw, _packet())

    assert "企业资料检索结果" not in result.content
    assert "tool_name" not in result.content
    assert result.content.count(repeated) == 1
    assert "[来源 1]" in result.content
    assert result.source_notes[0]["title"] == "FD-F 产品手册"
    assert result.duplicate_paragraph_count == 1


def test_professional_renderers_share_clean_canonical_content():
    artifact = {
        "title": "FD-F 多功能食品安全检测仪升级方案",
        "artifact_code": "ART-20260722-TEST",
        "artifact_label": "客户解决方案",
        "version_number": 2,
        "approval_status": "approved",
        "quality_score": 93,
        "content_markdown": (
            "# FD-F 多功能食品安全检测仪升级方案\n\n"
            "## 项目背景与客户目标\n企业需要完成设备升级。[EVID:doc-1:chunk-1]\n\n"
            "## 推荐配置\n| 模块 | 说明 |\n| --- | --- |\n| 主机 | 待核验 |"
        ),
    }

    docx = render_artifact_docx(artifact, _packet(), {"name": "飞达科技"})
    pdf = render_artifact_pdf(artifact, _packet(), {"name": "飞达科技"})
    xlsx = render_artifact_xlsx(artifact, _packet())

    document = Document(io.BytesIO(docx))
    all_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "资料来源" in all_text
    assert "[来源 1]" in all_text
    assert "EVID:" not in all_text
    assert all_text.count(artifact["title"]) == 1
    assert len(document.tables) >= 3
    assert document.styles["Title"].font.size.pt == 28
    assert pdf.startswith(b"%PDF")
    assert xlsx.startswith(b"PK")


def test_quality_gate_blocks_internal_trace_and_unsupported_numbers():
    spec = enrich_artifact_spec(
        ArtifactSpec(
            artifact_type=ArtifactType.CUSTOMER_SOLUTION,
            audience=ArtifactAudience.CUSTOMER,
            strict_quality=True,
            external_delivery=True,
        )
    )
    text = "\n\n".join(
        f"## {title}\ntool_result: 原始结果，保证 100% 达标。"
        for title in spec.required_sections
    )

    quality = evaluate_text_artifact(text, spec, _packet())
    codes = {item["code"] for item in quality["findings"]}

    assert contains_internal_trace_markers(text)
    assert quality["ready"] is False
    assert {"internal_trace_leakage", "numeric_claims_unsupported"} <= codes


def test_delivery_golden_set_covers_five_instrument_families():
    dataset = json.loads(
        (
            Path(__file__).resolve().parents[2]
            / "evals"
            / "datasets"
            / "artifact_delivery_golden.json"
        ).read_text(encoding="utf-8")
    )

    assert len(dataset["cases"]) == 5
    for case in dataset["cases"]:
        spec = infer_artifact_spec(case["request"])
        assert spec.artifact_type.value == case["artifact_type"]
        assert spec.instrument_line == case["instrument_line"]


def test_explicit_three_thousand_character_request_becomes_a_hard_contract():
    spec = infer_artifact_spec("参考企业资料生成不少于3000字的正式 Word 客户方案")

    assert spec.artifact_type == ArtifactType.CUSTOMER_SOLUTION
    assert spec.target_character_count == 3000
    assert spec.minimum_character_count == 3000
    assert spec.minimum_table_count == 3


def test_quality_gate_rejects_the_observed_loadknowledge_outline_failure():
    spec = enrich_artifact_spec(
        ArtifactSpec(
            artifact_type=ArtifactType.CUSTOMER_SOLUTION,
            audience=ArtifactAudience.CUSTOMER,
            strict_quality=True,
            external_delivery=True,
        )
    )
    failed_output = (
        "# 【loadknowledge】\n\n"
        "[企业资料检索结果]\n\n"
        "## 项目背景与客户目标\n当前资料不足。\n\n"
        "## 推荐配置与选型依据\n建议重新上传资料。"
    )

    quality = evaluate_text_artifact(failed_output, spec, _packet())
    codes = {item["code"] for item in quality["findings"]}

    assert quality["ready"] is False
    assert {
        "title_missing_or_generic",
        "required_sections_missing",
        "content_too_short",
        "structured_components_missing",
    } <= codes


class _FakeQuery:
    def __init__(self, table: str, rows: dict[str, list[dict]]):
        self.table = table
        self.rows = rows
        self.pending = None

    def select(self, *_args, **_kwargs):
        return self

    def eq(self, *_args, **_kwargs):
        return self

    def in_(self, *_args, **_kwargs):
        return self

    def ilike(self, *_args, **_kwargs):
        return self

    def limit(self, *_args, **_kwargs):
        return self

    def insert(self, value):
        self.pending = value
        return self

    async def execute(self):
        if self.pending is not None:
            values = self.pending if isinstance(self.pending, list) else [self.pending]
            self.rows.setdefault(self.table, []).extend(values)
            return SimpleNamespace(data=values)
        return SimpleNamespace(data=self.rows.get(self.table, []))


class _FakeDB:
    def __init__(self, rows=None):
        self.rows = rows or {}

    def table(self, name):
        return _FakeQuery(name, self.rows)


@pytest.mark.asyncio
async def test_explicit_enterprise_document_is_split_into_citable_topics(monkeypatch):
    async def empty_retrieval(**_kwargs):
        return EvidencePacket(topics=[], covered_topics=[], missing_topics=[])

    monkeypatch.setattr(
        "app.services.artifact_evidence_compiler.retrieve_agent_evidence",
        empty_retrieval,
    )
    topic_sections = (
        "一、客户行业场景与样品\n客户面向食品安全监管场景，检测对象覆盖现场样品与实验室样品。",
        "二、产品型号参数和检测能力\nFD-F 型号的参数、量程、精度、检出限与检测性能已经核验。",
        "三、适用标准政策\n方案适用的国家标准、行业标准、政策法规与实施规范已经归档。",
        "四、竞品参数\n同类竞品型号、公开参数、性能差异与国产替代边界已经整理。",
        "五、授权客户案例\n已授权客户案例包含项目部署、应用过程、验收结果与用户反馈。",
        "六、安装培训维保条款\n安装、培训、校准、维保、保修与售后响应条款已经确认。",
    )
    long_sections = "\n".join(
        section + (" 本节资料已经企业审核，可用于方案分析与引用。" * 30)
        for section in topic_sections
    )
    db = _FakeDB(
        {
            "documents": [
                {
                    "id": "11111111-1111-4111-8111-111111111111",
                    "name": "FD-F 完整产品方案.docx",
                    "doc_type": "product",
                    "review_status": "approved",
                    "source_version": "2026.2",
                    "extracted_data": {"full_text_context": long_sections},
                }
            ]
        }
    )
    spec = enrich_artifact_spec(
        ArtifactSpec(
            artifact_type=ArtifactType.CUSTOMER_SOLUTION,
            strict_quality=True,
            external_delivery=True,
        )
    )

    packet = await compile_artifact_evidence(
        query="参考 FD-F 完整产品方案.docx 生成食品安全检测仪升级方案",
        spec=spec,
        organization_id="org-1",
        user_id="user-1",
        db=db,
        selected_document_ids=[],
    )

    assert len(packet.records) >= 6
    assert packet.coverage == 1.0
    assert packet.sufficient is True
    assert "FD-F 完整产品方案.docx" in packet.prompt_context


@pytest.mark.asyncio
async def test_generation_pipeline_persists_version_evidence_and_quality(monkeypatch):
    spec = enrich_artifact_spec(
        ArtifactSpec(
            artifact_type=ArtifactType.CUSTOMER_SOLUTION,
            strict_quality=True,
            external_delivery=True,
        )
    )
    records = [
        EvidenceRecord(
            document_id=f"doc-{index}",
            chunk_id=f"chunk-{index}",
            title=f"企业资料 {index}",
            excerpt="已核验的产品、交付与服务事实。",
            score=1.0,
            source_version="2026.2",
            purposes=[topic],
        )
        for index, topic in enumerate(spec.retrieval_topics, 1)
    ]
    packet = EvidencePacket(
        records=records,
        topics=spec.retrieval_topics,
        covered_topics=spec.retrieval_topics,
        coverage=1.0,
        minimum_record_count=3,
        sufficient=True,
        prompt_context="\n".join(
            f"[{item.citation_id}] {item.excerpt}" for item in records
        ),
        fingerprint="golden-evidence",
    )

    async def fake_compile(**_kwargs):
        return packet

    sections = []
    for index, title in enumerate(spec.required_sections):
        body = (
            f"{title}围绕客户目标、企业能力和实施边界展开。"
            "本章先归纳企业资料中的已核验事实，再结合客户应用场景分析其适用性，"
            "明确推荐理由、限制条件、验收口径和责任边界。对于资料未覆盖的参数、"
            "价格、交期与承诺，统一保留待核验状态，不将推断写成事实。"
            "实施时应由售前、技术、交付与客户负责人共同确认输入条件，形成可追溯的"
            "记录，并在进入下一阶段前完成证据复核。最终建议以客户价值为导向，"
            "优先解决当前业务痛点，同时保留后续扩展空间和明确的下一步行动。"
            "项目推进过程中还应建立需求确认、配置冻结、到货检查、安装调试、培训考核"
            "和验收归档的阶段门，每个阶段记录输入、负责人、输出与异常处置方式，"
            "确保方案不仅能够阅读，也能够直接转化为可执行、可检查、可复盘的工作计划。"
        )
        if index < 3:
            body += (
                "\n\n| 核验维度 | 当前结论 | 证据状态 | 下一步 |\n"
                "| --- | --- | --- | --- |\n"
                "| 客户需求 | 已完成场景归纳 | 已核验 | 确认验收口径 |\n"
                "| 企业能力 | 具备对应方案基础 | 已核验 | 完成配置复核 |"
            )
        sections.append(
            {
                "title": title,
                "content": body,
                "evidence_refs": [records[index % len(records)].citation_id],
            }
        )

    async def fake_chat(**_kwargs):
        return SimpleNamespace(
            content=json.dumps(
                {
                    "title": "食品安全检测仪升级换代整体方案",
                    "executive_summary": (
                        "本方案面向食品安全检测能力升级与持续运营需求，基于企业已授权的"
                        "产品、政策、竞品、案例和服务资料，形成从需求确认、配置选型、"
                        "实施验收到售后保障的完整路径。方案强调事实可追溯、参数可核验、"
                        "责任边界清晰，并将缺少证据的商务与技术事项保留为人工复核项，"
                        "便于客户快速判断方案价值并推进下一步技术交流。"
                    ),
                    "sections": sections,
                    "verification_items": ["最终报价由负责人确认"],
                },
                ensure_ascii=False,
            ),
            finish_reason="stop",
            model_code="deepseek-v4-flash",
            usage={"total_tokens": 1200},
        )

    async def ignore_quality_event(**_kwargs):
        return None

    monkeypatch.setattr(
        "app.services.artifact_generation_service.compile_artifact_evidence",
        fake_compile,
    )
    monkeypatch.setattr(
        "app.services.artifact_generation_service.llm_gateway.chat", fake_chat
    )
    monkeypatch.setattr(
        "app.services.artifact_generation_service.persist_artifact_quality_event",
        ignore_quality_event,
    )
    db = _FakeDB()

    result = await generate_artifact(
        db=db,
        organization_id="11111111-1111-4111-8111-111111111111",
        user_id="22222222-2222-4222-8222-222222222222",
        original_request="给客户生成食品安全检测仪升级解决方案",
        source_content="已有聊天草稿",
        artifact_type=ArtifactType.CUSTOMER_SOLUTION,
        audience=ArtifactAudience.CUSTOMER,
        target_character_count=3000,
        review_confirmed=True,
    )

    assert result["quality"]["ready"] is True, result["quality"]
    assert result["quality"]["metrics"]["character_count"] >= 3000
    assert result["quality"]["metrics"]["table_count"] >= 3
    assert result["approval_status"] == "approved"
    assert len(db.rows["artifacts"]) == 1
    assert len(db.rows["artifact_versions"]) == 1
    assert len(db.rows["artifact_evidence_links"]) == len(records)
