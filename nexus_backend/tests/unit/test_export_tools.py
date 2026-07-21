import base64
import io

import pytest
from docx import Document

from app.tools.export_tools import export_to_docx


@pytest.mark.asyncio
async def test_export_to_docx_preserves_headings_lists_and_tables():
    result = await export_to_docx.ainvoke(
        {
            "content": "\n".join(
                [
                    "## 推荐配置",
                    "- 液相色谱主机",
                    "- 紫外检测器",
                    "",
                    "| 型号 | 数量 |",
                    "| --- | ---: |",
                    "| LC-100 | 1 |",
                ]
            ),
            "filename": "客户方案",
            "title": "制药实验室解决方案",
            "format_type": "markdown",
        }
    )

    assert result["success"] is True
    assert result["filename"] == "客户方案.docx"

    document = Document(io.BytesIO(base64.b64decode(result["content_base64"])))
    paragraph_text = [paragraph.text for paragraph in document.paragraphs]
    assert "制药实验室解决方案" in paragraph_text
    assert "推荐配置" in paragraph_text
    assert "液相色谱主机" in paragraph_text
    assert document.tables[0].cell(1, 0).text == "LC-100"

